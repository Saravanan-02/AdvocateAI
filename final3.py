import os
import streamlit as st
import requests
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer, CrossEncoder
import tempfile
import pickle
import random
import re
from rank_bm25 import BM25Okapi
import nltk
from nltk.corpus import stopwords
import string
import base64
import json
from datetime import datetime
from io import BytesIO

# --- PDF Dependencies ---
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import inch

# --- WORD (.docx) Dependencies ---
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("⚠️ Word export unavailable. Install 'python-docx' to enable this feature.", icon="📄")

# --- PDF Text Extraction Dependencies ---
try:
    import PyPDF2
    from PyPDF2.errors import PdfReadError
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    st.warning("⚠️ PDF upload unavailable. Install 'PyPDF2' to enable this feature.", icon="📄")


# ==========================
# NLTK DATA DOWNLOADS
# ==========================
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

# ==========================
# CONFIG & API KEY
# ==========================

# --- STREAMLIT SECRETS CONFIGURATION ---
OPENROUTER_API_KEY = None
try:
    OPENROUTER_API_KEY = st.secrets.get("OPENROUTER_API_KEY")
except:
    pass

if not OPENROUTER_API_KEY:
    OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")

BASE_URL = "https://openrouter.ai/api/v1"
INDEX_FILE = "faiss_advanced_index.bin"
METADATA_FILE = "metadata.pkl"

# Check if API key is configured
if not OPENROUTER_API_KEY:
    st.error("⚠️ OPENROUTER_API_KEY not found. Please configure it in Streamlit secrets.")
    st.info("Add your API key to `.streamlit/secrets.toml` with: `OPENROUTER_API_KEY = \"your-key-here\"`")
    st.stop()

# ==========================
# JAVASCRIPT FOR NATIVE TTS
# ==========================
def inject_tts_javascript():
    js_code = """
    function speakText(text) {
        if ('speechSynthesis' in window) {
            let cleanText = text.replace(/<[^>]*>/g, '');
            cleanText = cleanText.replace(/\\*\\*(.*?)\\*\\*/g, '$1');
            cleanText = cleanText.replace(/\\n/g, ' ');
            if (window.speechSynthesis.speaking) {
                window.speechSynthesis.cancel();
            }
            const utterance = new SpeechSynthesisUtterance(cleanText);
            utterance.rate = 1.0; 
            utterance.pitch = 1.0; 
            utterance.volume = 1.0;
            utterance.lang = 'en-US'; 
            window.speechSynthesis.speak(utterance);
        } else {
            alert("Speech Synthesis not supported in this browser.");
        }
    }
    """
    return f"<script>{js_code}</script>"

# ==========================
# CACHED RESOURCE LOADING
# ==========================
@st.cache_resource
def load_models_and_index():
    embedding_model = SentenceTransformer("all-MiniLM-L6-v2", device="cpu") 
    reranker_model = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2", device="cpu")
    try:
        index = faiss.read_index(INDEX_FILE)
        with open(METADATA_FILE, "rb") as f:
            metadata = pickle.load(f)
        return embedding_model, reranker_model, index, metadata
    except FileNotFoundError:
        return embedding_model, reranker_model, None, None

# ==========================
# HELPER FUNCTIONS 
# ==========================
def create_extractive_summary(text, max_sentences=3):
    if not text:
        return ""
    try:
        sentences = nltk.sent_tokenize(text)
        summary = " ".join(sentences[:max_sentences])
        return summary
    except LookupError:
        sentences = re.split(r'(?<=[.!?]) +', text.strip())
        summary = " ".join(sentences[:max_sentences])
        return summary

def truncate_text(text, max_words=500):
    words = text.split()
    return " ".join(words[:max_words])

def summarize_source_with_llm(source_text, model="openai/gpt-5-mini"):
    return create_extractive_summary(source_text, max_sentences=3) 

def preprocess_text(text):
    if not text:
        return []
    text = text.lower()
    text = ''.join([c for c in text if c not in string.punctuation])
    tokens = text.split()
    stop_words = set(stopwords.words('english'))
    return [w for w in tokens if w not in stop_words]

def bm25_search(query, metadata, top_k=10):
    tokenized_corpus = [preprocess_text(m['text']) for m in metadata]
    if not tokenized_corpus:
        return []
    bm25 = BM25Okapi(tokenized_corpus)
    tokenized_query = preprocess_text(query)
    if not tokenized_query:
        return []
    scores = bm25.get_scores(tokenized_query)
    ranked_indices = np.argsort(scores)[::-1]
    results = []
    for idx in ranked_indices[:top_k]:
        results.append({
            'source': metadata[idx].get('source', 'BM25 Result'),
            'text': metadata[idx]['text'],
            'score': scores[idx]
        })
    return results

def reciprocal_rank_fusion(results_list, k=60):
    fused_ranks = {}
    for results in results_list:
        if not results:
            continue
        for rank, item in enumerate(results):
            doc_id = item['source'] + '_' + item['text']
            fused_ranks[doc_id] = fused_ranks.get(doc_id, 0) + 1 / (k + rank)
    sorted_items = sorted(fused_ranks.items(), key=lambda x: x[1], reverse=True)
    unique_items = {item['source'] + '_' + item['text']: item for results in results_list if results for item in
                    results}
    fused_results = [unique_items[doc_id] for doc_id, _ in sorted_items if doc_id in unique_items]
    return fused_results

def retrieve_and_rerank(query, index, metadata, embed_model, reranker_model, top_k=2):
    """Fixed version with proper FAISS handling and metadata validation"""
    try:
        # Validate metadata first - SILENTLY
        if not metadata or not isinstance(metadata, list) or len(metadata) == 0:
            # Try BM25 fallback without error message
            try:
                bm25_results = bm25_search(query, metadata if metadata else [], top_k=top_k)
                if bm25_results:
                    top_chunks = [truncate_text(r["text"], max_words=200) for r in bm25_results]
                    return "\n\n".join(top_chunks), bm25_results
            except:
                pass
            return "", []
        
        query_emb = embed_model.encode([query], convert_to_numpy=True)
        
        # Safe search with proper error handling
        metadata_size = len(metadata)
        search_k = min(top_k * 4, index.ntotal if index else 0, metadata_size)
        
        if search_k == 0 or not index:
            # Fallback to BM25 only
            bm25_results = bm25_search(query, metadata, top_k=top_k)
            if bm25_results:
                top_chunks = [truncate_text(r["text"], max_words=200) for r in bm25_results]
                return "\n\n".join(top_chunks), bm25_results
            return "", []
        
        distances, indices = index.search(query_emb, search_k)
        
        # Filter valid indices with strict bounds checking
        semantic_results = []
        for idx in indices[0]:
            idx = int(idx)  # Ensure it's a Python int
            if idx >= 0 and idx < metadata_size:
                try:
                    semantic_results.append(metadata[idx])
                except (IndexError, KeyError):
                    continue
        
        # BM25 search as fallback
        bm25_results = bm25_search(query, metadata, top_k=search_k)
        
        # Fusion
        fused_candidates = reciprocal_rank_fusion([semantic_results, bm25_results])
        
        if not fused_candidates:
            # Fallback to just BM25 if fusion fails
            fused_candidates = bm25_results[:top_k] if bm25_results else []
        
        if not fused_candidates:
            return "", []
        
        pairs = [[query, c["text"]] for c in fused_candidates if "text" in c]
        if not pairs:
            return "", []
        
        scores = reranker_model.predict(pairs)
        reranked = sorted(zip(fused_candidates, scores), key=lambda x: x[1], reverse=True)[:top_k]
        top_chunks = [truncate_text(item[0]["text"], max_words=200) for item in reranked]
        sources = [{"text": item[0]["text"], "source": item[0].get("source", "Unknown")} for item in reranked]
        return "\n\n".join(top_chunks), sources
    except Exception as e:
        # Silent fallback - just use BM25
        try:
            bm25_results = bm25_search(query, metadata if metadata else [], top_k=top_k)
            if bm25_results:
                top_chunks = [truncate_text(r["text"], max_words=200) for r in bm25_results]
                return "\n\n".join(top_chunks), bm25_results
        except:
            pass
        return "", []

def dynamic_query_generator(question):
    """Generates three distinct prompt variations for legal research."""
    base_prompt = question.strip()
    q_lower = base_prompt.lower()
    variations = [base_prompt]
    topic = re.sub(r'^(tell me about|what is|define|explain|how to)\s+', '', q_lower, flags=re.IGNORECASE).strip()
    if q_lower.startswith("how to"):
        refined_prompt_2 = f"What are the steps, legal guidelines, and mandatory procedure required for {topic}?"
    elif topic and len(topic) < len(q_lower):
        refined_prompt_2 = f"Provide a complete definition, key elements, and relevant section references for the legal concept: {topic}."
    else:
        refined_prompt_2 = f"Conduct a formal structural analysis of the legal principle or section: {base_prompt}."
    if re.search(r'ipc|act|section|law', q_lower):
        refined_prompt_3 = f"Cite recent Supreme Court case law, mandatory punishment/penalty, and any recent amendments related to: {base_prompt}."
    else:
        refined_prompt_3 = f"Compare {base_prompt} with related legal concepts and explain its application in a modern context, referencing relevant precedents."
    if refined_prompt_2 not in variations: variations.append(refined_prompt_2)
    if refined_prompt_3 not in variations and len(variations) < 3: variations.append(refined_prompt_3)
    if len(variations) < 3:
        variations.append(f"What are the common challenges or defenses related to {base_prompt}?")
    
    return variations[:3]

@st.cache_data(ttl=600)
def extract_text_from_pdfs(uploaded_files):
    if not uploaded_files or not PYPDF2_AVAILABLE:
        return ""
    all_text = []
    with st.spinner(f"Processing {len(uploaded_files)} PDF(s)..."):
        for file in uploaded_files:
            try:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        all_text.append(text)
            except PdfReadError:
                st.warning(f"Could not read '{file.name}'. The file might be corrupt or encrypted.")
            except Exception as e:
                st.error(f"An error occurred while processing '{file.name}': {e}")
            file.seek(0)
    return "\n\n".join(all_text)

# ==========================
# PROMPT CONSTRUCTION
# ==========================
def construct_final_prompt(question, rag_summary, uploaded_context_summary):
    rag_lines = rag_summary.strip().splitlines()
    if not rag_lines:
        knowledge_block = "No relevant knowledge found in the database."
    else:
        knowledge_block = "\n".join(rag_lines[:3]) 
        
    if not uploaded_context_summary:
        uploaded_block = "No context provided from uploaded documents."
    else:
        uploaded_block = truncate_text(uploaded_context_summary, max_words=1500)

    final_prompt = f"""
    Advocate AI: Answer the following question clearly, concisely, and formally.
    
    **Instructions:**
    1. Provide a professional, structured legal response based on the **Selected Question** below.
    2. **Crucially, maintain clear line breaks and use markdown (like bullet points or list numbering) for structure.**
    3. Do NOT include any initial greetings, intros, or headings. Start directly with the answer content.
    4. You have two sources of information. Prioritize context from 'Uploaded Docs' if it is relevant.
    5. Cite references clearly, e.g., (Source: SourceFileName.pdf) or (Source: Uploaded Document).
    
    **Selected Question:** {question}
    
    ---
    **Knowledge (From Database):**
    {knowledge_block}
    
    ---
    **Knowledge (From Uploaded Docs):**
    {uploaded_block}
    """
    return final_prompt.strip()


# ==========================
# LLM API CALL (STREAMING)
# ==========================
def call_openrouter(model, prompt, sources=None):
    headers = {"Authorization": f"Bearer {OPENROUTER_API_KEY}", "Content-Type": "application/json"}
    # Add "stream": True to enable streaming
    data = {"model": model, "messages": [{"role": "user", "content": prompt}], "stream": True}
    
    try:
        # Add stream=True to the requests call
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions", 
            headers=headers, 
            json=data, 
            timeout=30, 
            stream=True
        )
        response.raise_for_status()
        
        # Iterate over the streaming response
        for line in response.iter_lines():
            if line:
                decoded_line = line.decode('utf-8')
                if decoded_line.startswith('data: '):
                    json_data = decoded_line[6:]
                    if json_data == "[DONE]":
                        break
                    try:
                        chunk = json.loads(json_data)
                        if "choices" in chunk and chunk["choices"][0].get("delta", {}).get("content"):
                            # Yield the text chunk
                            yield chunk["choices"][0]["delta"]["content"]
                    except json.JSONDecodeError:
                        continue # Ignore empty or malformed lines
    
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 401:
            err_msg = f"**LLM Error (401 Unauthorized):** Cannot connect. Please fix your API Key. {e}"
            st.error(err_msg)
            yield err_msg
        else:
            err_msg = f"HTTP Error connecting to LLM: {e}"
            st.error(err_msg)
            yield err_msg
    except requests.exceptions.Timeout:
        err_msg = "Error: Request timed out."
        st.error(err_msg)
        yield err_msg
    except Exception as e:
        err_msg = f"Error connecting to LLM: {e}"
        st.error(err_msg)
        yield err_msg

def generate_suggested_prompts(last_question, last_answer):
    """Generates 3 follow-up prompts."""
    suggestion_prompt = f"""
    Based on the last legal question and its answer, generate 3 brief, distinct, and relevant follow-up questions a legal professional might ask next.
    
    **Rules:**
    1. Return ONLY the 3 questions.
    2. Separate each question with a newline.
    3. Do NOT number them or use bullet points.
    4. The questions should be specific and insightful.
    
    **Last Question:**
    {last_question}
    
    **Last Answer:**
    {truncate_text(last_answer, 300)}
    """
    # Note: This call is still non-streaming, which is fine for short suggestions.
    # We need to collect the stream into a full answer
    stream = call_openrouter("openai/gpt-5-mini", suggestion_prompt)
    answer = "".join(stream)
    
    if "Error" in answer:
        return []
    prompts = [
        p.strip() for p in answer.split('\n') 
        if p.strip() and p.strip().endswith('?')
    ]
    return prompts[:3]


# =========================================
# KEYWORD HIGHLIGHTING
# =========================================
def highlight_with_style(text, search_terms_raw):
    """
    Highlights terms and returns text, plus a boolean if anything was found.
    """
    found_keyword = False
    if not search_terms_raw:
        return text, found_keyword
    search_terms = [re.escape(term.strip()) for term in search_terms_raw.split(',') if term.strip()]
    if not search_terms:
        return text, found_keyword
        
    pattern = re.compile(r'(' + '|'.join(search_terms) + r')', re.IGNORECASE)
    
    highlight_style = 'background-color: #ADD8E6; color: black; border-radius: 3px; padding: 2px 4px; display: inline; font-weight: bold;'
    
    def replacer(match):
        nonlocal found_keyword
        found_keyword = True
        return f'<span style="{highlight_style}">{match.group(0)}</span>'
        
    highlighted_text = pattern.sub(replacer, text)
    return highlighted_text, found_keyword

def check_keywords_in_chat(messages, keyword_terms):
    """
    Loops through chat history to check if keywords exist.
    Returns True if ANY keyword is found in ANY message.
    """
    if not keyword_terms or not keyword_terms.strip():
        return True  # No keywords to search for
    
    # Split keywords and clean them
    keywords = [term.strip().lower() for term in keyword_terms.split(',') if term.strip()]
    
    if not keywords:
        return True
    
    # Search through all messages
    for message in messages:
        content = message.get("content", "").lower()
        for keyword in keywords:
            if keyword in content:
                return True
    
    return False


# ===================================================================
# DOWNLOAD CHAT CONTENT (PDF & WORD)
# ===================================================================

def clean_markdown_and_linebreaks(text):
    text = re.sub(r'#+\s*', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text, flags=re.DOTALL)
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text, flags=re.DOTALL)
    text = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', text)
    text = re.sub(r'^[\s]*[\*-]\s+', '&bull; ', text, flags=re.MULTILINE)
    text = re.sub(r'\n\n', '<br/><br/>', text)
    text = re.sub(r'\n', '<br/>', text)
    return text.strip()

def clean_text_for_docx(text):
    text = re.sub(r'#+\s*', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'\*(.*?)\*', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', text)
    text = re.sub(r'^[\s]*[\*-]\s+', '\u2022 ', text, flags=re.MULTILINE)
    text = re.sub(r'<[^>]*>', '', text)
    return text.strip()

def generate_qna_content_pdf(messages, chat_name):
    buffer = BytesIO()
    margin = 0.75 * inch 
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            leftMargin=margin, rightMargin=margin,
                            topMargin=margin, bottomMargin=margin)
    styles = getSampleStyleSheet()
    story = []
    line_spacing = 1.5
    font_size_main = 15
    font_size_source = 12
    styles.add(ParagraphStyle(name='Question', fontName='Times-Bold', fontSize=font_size_main, leading=font_size_main * line_spacing, spaceAfter=0, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name='AnswerHTML', fontName='Times-Roman', fontSize=font_size_main, leading=font_size_main * line_spacing, spaceAfter=8, leftIndent=10))
    styles.add(ParagraphStyle(name='SourceHeader', fontName='Times-Bold', fontSize=font_size_source, leading=font_size_source * line_spacing, spaceBefore=5, spaceAfter=2, leftIndent=10))
    styles.add(ParagraphStyle(name='SourceItem', fontName='Times-Roman', fontSize=font_size_source, leading=font_size_source * line_spacing, spaceAfter=2, leftIndent=20))
    styles.add(ParagraphStyle(name='ChatTitle', fontName='Times-Bold', fontSize=18, spaceAfter=10, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name='SubTitle', fontName='Times-Bold', fontSize=15, spaceAfter=5, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name='Date', fontName='Times-Roman', fontSize=12, spaceAfter=20, alignment=TA_LEFT))

    story.append(Paragraph("ADVOCATE AI CHAT SESSION", styles['ChatTitle']))
    story.append(Paragraph(f"Topic: {chat_name}", styles['SubTitle']))
    story.append(Paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Date']))
    story.append(Spacer(1, 0.25 * inch))

    q_count = 1
    for i, message in enumerate(messages):
        if message["role"] == "user":
            if i + 1 < len(messages) and messages[i+1]["role"] == "assistant":
                answer_message = messages[i+1]
                question_text = f"Q{q_count}: {clean_markdown_and_linebreaks(message['content'])}"
                story.append(Paragraph(question_text, styles['Question']))
                story.append(Spacer(1, 15))
                html_answer = clean_markdown_and_linebreaks(answer_message['content'])
                story.append(Paragraph(html_answer, styles['AnswerHTML']))
                if answer_message.get("sources"):
                    story.append(Paragraph("--- Sources Cited (from Database) ---", styles['SourceHeader']))
                    for j, source in enumerate(answer_message["sources"]):
                        source_summary = summarize_source_with_llm(source['text']) 
                        source_line = (f"&bull; <b>Source {j+1}:</b> {source['source']}<br/>"
                                       f"&nbsp;&nbsp;&nbsp;&nbsp;<i>Snippet:</i> {source_summary}")
                        story.append(Paragraph(source_line, styles['SourceItem']))
                story.append(Spacer(1, 0.2 * inch))
                q_count += 1
    try:
        doc.build(story)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"Error generating PDF: {e}")
        return b""

def generate_qna_content_word(messages, chat_name):
    if not DOCX_AVAILABLE:
        return b""
    
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(15)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    title = document.add_heading("ADVOCATE AI CHAT SESSION", level=1)
    title.runs[0].font.name = 'Times New Roman'
    topic = document.add_paragraph()
    topic.add_run(f"Topic: {chat_name}").bold = True
    document.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").runs[0].font.size = Pt(12)
    document.add_paragraph()
    q_count = 1
    for i, message in enumerate(messages):
        if message["role"] == "user":
            if i + 1 < len(messages) and messages[i+1]["role"] == "assistant":
                answer_message = messages[i+1]
                q_para = document.add_paragraph()
                q_run = q_para.add_run(f"Q{q_count}: {clean_text_for_docx(message['content'])}")
                q_run.bold = True
                document.add_paragraph()
                a_para = document.add_paragraph(clean_text_for_docx(answer_message['content']))
                a_para.paragraph_format.left_indent = Inches(0.25)
                if answer_message.get("sources"):
                    s_header_para = document.add_paragraph()
                    s_header_para.paragraph_format.left_indent = Inches(0.25)
                    s_run = s_header_para.add_run("--- Sources Cited (from Database) ---")
                    s_run.font.size = Pt(12)
                    s_run.italic = True
                    for j, source in enumerate(answer_message["sources"]):
                        source_summary = summarize_source_with_llm(source['text'])
                        source_text = f"\u2022 Source {j+1}: {source['source']}\n\tSnippet: {source_summary}"
                        src_para = document.add_paragraph(source_text)
                        src_para.paragraph_format.left_indent = Inches(0.5)
                        for run in src_para.runs:
                            run.font.size = Pt(12)
                document.add_paragraph()
                q_count += 1
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

# ==========================
# STREAMLIT APP SETUP
# ==========================
st.set_page_config(page_title="Advocate AI Optimized", layout="wide")
st.markdown(inject_tts_javascript(), unsafe_allow_html=True) 
embed_model, reranker_model, folder_index, folder_metadata = load_models_and_index()

if folder_index is None:
    st.error("Knowledge base not found. Run build_index_advanced.py.")
    st.stop()

# Initialize session state 
if "messages" not in st.session_state: st.session_state.messages = []
if "uploaded_pdfs_data" not in st.session_state: st.session_state.uploaded_pdfs_data = None
if "chat_sessions" not in st.session_state: st.session_state.chat_sessions = {}
if "active_chat" not in st.session_state: st.session_state.active_chat = None
if "selected_model" not in st.session_state: st.session_state.selected_model = "openai/gpt-5"
if "keyword_found" not in st.session_state: st.session_state.keyword_found = True
if "suggested_prompts" not in st.session_state: st.session_state.suggested_prompts = []
if "pending_prompts" not in st.session_state: st.session_state.pending_prompts = []
if "processing" not in st.session_state: st.session_state.processing = False
if "active_stream" not in st.session_state: st.session_state.active_stream = None
if "stream_sources" not in st.session_state: st.session_state.stream_sources = []
if "last_user_prompt" not in st.session_state: st.session_state.last_user_prompt = ""


# Inject Custom CSS
custom_css = f"""
<style>
/* Base UI Cleanup */
.main {{ 
    padding-top: 20px; 
    background: linear-gradient(to bottom, #f8f9fa 0%, #ffffff 100%);
}}

/* Hide Streamlit branding for cleaner look */
#MainMenu {{visibility: hidden;}}
footer {{visibility: hidden;}}

/* Professional Prompt Card Styling */
.prompt-selection-container {{
    background: white;
    border-radius: 16px;
    padding: 30px;
    box-shadow: 0 8px 24px rgba(0, 0, 0, 0.08);
    margin: 20px 0;
}}

.prompt-card {{
    background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%);
    border: 2px solid #e8eaed;
    border-radius: 16px;
    padding: 24px;
    margin: 15px 0;
    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.04);
    transition: all 0.35s cubic-bezier(0.4, 0, 0.2, 1);
    cursor: pointer;
    position: relative;
    overflow: hidden;
}}

.prompt-card::before {{
    content: '';
    position: absolute;
    top: 0;
    left: 0;
    width: 100%;
    height: 4px;
    background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
    transform: scaleX(0);
    transition: transform 0.35s ease;
}}

.prompt-card:hover {{
    border-color: #1e3c72;
    box-shadow: 0 8px 24px rgba(30, 60, 114, 0.12);
    transform: translateY(-4px);
}}

.prompt-card:hover::before {{
    transform: scaleX(1);
}}

.prompt-card.original {{
    border-left: 5px solid #d4af37;
    background: linear-gradient(145deg, #fffef7 0%, #faf8f0 100%);
}}

.prompt-card.original::before {{
    background: linear-gradient(90deg, #d4af37 0%, #f4d03f 100%);
}}

.prompt-card.refined {{
    border-left: 5px solid #1e3c72;
}}

.prompt-label {{
    font-size: 12px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    color: #5f6368;
    margin-bottom: 12px;
    display: flex;
    align-items: center;
    gap: 8px;
}}

.prompt-label .icon {{
    font-size: 16px;
}}

.prompt-text {{
    font-size: 15px;
    color: #202124;
    line-height: 1.7;
    font-weight: 500;
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
}}

/* Section Headers */
.section-header {{
    font-size: 20px;
    font-weight: 700;
    color: #1e3c72;
    margin: 30px 0 20px 0;
    padding-bottom: 10px;
    border-bottom: 3px solid #d4af37;
    display: flex;
    align-items: center;
    gap: 10px;
}}

/* Sidebar Chat Button Styling */
.stButton>button {{
    border-radius: 10px;
    transition: all 0.25s ease;
    margin-top: 5px;
    height: 42px;
    border: 1.5px solid #e0e0e0;
    font-weight: 500;
}}

.stButton>button:hover {{ 
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.12); 
    border: 1.5px solid #1e3c72;
    transform: translateX(2px);
}}

/* Primary button for Active Chat Session */
.stButton[data-testid*="primary"]>button {{
    background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
    color: white;
    border: 1.5px solid #1e3c72;
    box-shadow: 0 4px 8px rgba(30, 60, 114, 0.2);
}}

/* Chat Message Bubbles */
div[data-testid="stChatMessage"] {{
    padding: 16px;
    border-radius: 12px;
    margin-bottom: 16px;
    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05);
}}

/* User Message */
div[data-testid="stChatMessage"][data-testid*="user"] {{
    background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%);
    border-left: 4px solid #1e3c72;
}}

/* Assistant Message */
div[data-testid="stChatMessage"][data-testid*="assistant"] {{
    background: linear-gradient(135deg, #f5f5f5 0%, #eeeeee 100%);
    border-right: 4px solid #d4af37;
}}

/* TTS Button Styling */
.stButton button[kind="secondary"] {{
    background: linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%);
    color: #1e3c72;
    border: 1.5px solid #e0e0e0;
    font-size: 11px;
    padding: 4px 8px;
    height: 28px;
    margin-bottom: 10px;
    border-radius: 6px;
    transition: all 0.2s ease;
}}

.stButton button[kind="secondary"]:hover {{
    background: #1e3c72;
    color: white;
    border-color: #1e3c72;
}}

/* Suggestion Button Styling - FIXED */
.suggestion-container {{
    margin: 12px 0;
}}

.suggestion-btn {{
    background: white;
    border: 2px solid #e0e0e0;
    border-radius: 12px;
    padding: 18px 20px;
    transition: all 0.3s ease;
    text-align: left;
    font-size: 14px;
    line-height: 1.6;
    color: #202124;
    cursor: pointer;
    box-shadow: 0 2px 6px rgba(0, 0, 0, 0.04);
    display: block;
    width: 100%;
    word-wrap: break-word;
    overflow-wrap: break-word;
    white-space: normal;
    min-height: 60px;
}}

.suggestion-btn:hover {{
    background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%);
    border-color: #1e3c72;
    box-shadow: 0 4px 12px rgba(30, 60, 114, 0.15);
    transform: translateX(4px);
}}

/* Fix Streamlit button container for suggestions */
div[data-testid="column"] > div > div > div > button {{
    white-space: normal !important;
    height: auto !important;
    min-height: 60px !important;
    text-align: left !important;
    padding: 18px 20px !important;
    line-height: 1.6 !important;
    word-wrap: break-word !important;
}}

/* Ensure all buttons in main content area wrap properly */
.stButton > button {{
    white-space: normal !important;
    height: auto !important;
    padding: 12px 20px !important;
    line-height: 1.5 !important;
    word-break: break-word !important;
    text-align: left !important;
}}

/* Sidebar buttons keep their compact style */
[data-testid="stSidebar"] .stButton > button {{
    white-space: nowrap !important;
    height: 42px !important;
    text-align: center !important;
    overflow: hidden !important;
    text-overflow: ellipsis !important;
}}

/* Loading Animation */
@keyframes pulse {{
    0%, 100% {{ opacity: 1; transform: scale(1); }}
    50% {{ opacity: 0.7; transform: scale(0.98); }}
}}

.processing-indicator {{
    animation: pulse 2s ease-in-out infinite;
    background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
    color: white;
    padding: 20px;
    border-radius: 12px;
    text-align: center;
    font-weight: 600;
    font-size: 16px;
    box-shadow: 0 4px 16px rgba(30, 60, 114, 0.3);
    margin: 20px 0;
}}

.processing-subtext {{
    font-size: 13px;
    color: #e3f2fd;
    margin-top: 8px;
    font-weight: 400;
}}

/* Input Field Styling */
.stChatInput {{
    border-radius: 12px;
    border: 2px solid #e0e0e0;
    transition: all 0.3s ease;
}}

.stChatInput:focus-within {{
    border-color: #1e3c72;
    box-shadow: 0 0 0 3px rgba(30, 60, 114, 0.1);
}}

/* Expander Styling */
.streamlit-expanderHeader {{
    background: linear-gradient(135deg, #f8f9fa 0%, #ffffff 100%);
    border-radius: 8px;
    font-weight: 600;
    color: #1e3c72;
}}

/* Column Cards in Prompt Selection */
.stColumn {{
    padding: 8px;
}}
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)


# --- KEYWORD BUG FIX: Run check *before* sidebar is rendered ---
keyword_input = st.session_state.get("keyword_highlight_input", "")
if keyword_input and st.session_state.messages:
    st.session_state.keyword_found = check_keywords_in_chat(
        st.session_state.messages, 
        keyword_input
    )
else:
    st.session_state.keyword_found = True


# ==========================
# SIDEBAR 
# ==========================
with st.sidebar:
    
    st.header("🔑 Keyword Highlight")
    keyword_search_input = st.text_input("Enter keywords (comma-separated)", key="keyword_highlight_input", placeholder="e.g., defense, doctrine, contract")
    
    # Show keyword status only if there are messages and keywords entered
    if keyword_search_input and st.session_state.messages:
        if st.session_state.keyword_found:
            st.success("✅ Keywords found in chat", icon="✨")
        else:
            st.warning("⚠️ Keywords not found in current chat", icon="🔍")
    elif keyword_search_input and not st.session_state.messages:
        st.info("💬 Start a conversation to search for keywords", icon="ℹ️")
    
    st.markdown("---") 

    st.header("💬 Chat History")
    if st.button("➕ New Chat", use_container_width=True):
        chat_id = f"chat_{len(st.session_state.chat_sessions) + 1}_{random.randint(1000,9999)}"
        st.session_state.chat_sessions[chat_id] = {"name": "New Chat", "messages": [], "created_at": datetime.now()}
        st.session_state.active_chat = chat_id
        st.session_state.messages = []
        st.session_state.keyword_found = True
        st.session_state.suggested_prompts = []
        st.session_state.pending_prompts = []
        st.session_state.active_stream = None
        st.rerun()

    st.write("**Your Chats:**")
    if not st.session_state.chat_sessions:
        st.info("No chats yet.")
    else:
        sorted_chats = sorted(st.session_state.chat_sessions.items(), key=lambda x: x[1].get('created_at', datetime.min), reverse=True)
        for chat_id, session in sorted_chats:
            is_active = st.session_state.active_chat == chat_id
            chat_col1, chat_col2 = st.columns([4, 1])
            with chat_col1:
                button_type = "primary" if is_active else "secondary"
                if st.button(session["name"], key=f"chat_btn_{chat_id}", use_container_width=True, type=button_type):
                    st.session_state.active_chat = chat_id
                    st.session_state.messages = session["messages"]
                    st.session_state.keyword_found = True
                    st.session_state.suggested_prompts = []
                    st.session_state.pending_prompts = []
                    st.session_state.active_stream = None
                    st.rerun()
            with chat_col2:
                with st.popover("⋮"):
                    new_name = st.text_input("Rename chat", value=session["name"], key=f"rename_{chat_id}")
                    if st.button("Save", key=f"save_{chat_id}"):
                        st.session_state.chat_sessions[chat_id]["name"] = new_name
                        st.rerun()
                    if st.button("Delete", key=f"delete_{chat_id}"):
                        del st.session_state.chat_sessions[chat_id]
                        if st.session_state.active_chat == chat_id:
                            st.session_state.active_chat = None
                            st.session_state.messages = []
                            st.session_state.keyword_found = True
                            st.session_state.suggested_prompts = []
                            st.session_state.pending_prompts = []
                            st.session_state.active_stream = None
                        st.rerun()
    st.markdown("---")

    # --- DOWNLOAD & CLEAR HISTORY ---
    pdf_content_bytes = b''
    word_content_bytes = b''
    pdf_file_name = "AdvocateAI_Session.pdf"
    word_file_name = "AdvocateAI_Session.docx"
    current_messages = []
    current_chat_name = "Current_Unsaved_Chat"

    if st.session_state.active_chat and st.session_state.active_chat in st.session_state.chat_sessions:
        current_chat = st.session_state.chat_sessions[st.session_state.active_chat]
        current_chat_name = current_chat["name"]
        current_messages = current_chat["messages"]
    elif st.session_state.messages:
        current_messages = st.session_state.messages

    if current_messages:
        base_file_name = f"AdvocateAI_{current_chat_name.replace(' ', '_')[:20]}"
        pdf_file_name = f"{base_file_name}.pdf"
        word_file_name = f"{base_file_name}.docx"
        pdf_content_bytes = generate_qna_content_pdf(current_messages, current_chat_name)
        if DOCX_AVAILABLE:
            word_content_bytes = generate_qna_content_word(current_messages, current_chat_name)

    st.download_button(label="⬇️ Download Q&A as PDF", data=pdf_content_bytes, file_name=pdf_file_name, mime="application/pdf", use_container_width=True, disabled=(not pdf_content_bytes))
    
    if DOCX_AVAILABLE:
        st.download_button(label="⬇️ Download Q&A as Word", data=word_content_bytes, file_name=word_file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, disabled=(not word_content_bytes))
    
    if st.button("🗑️ Clear Chat History", use_container_width=True):
        st.session_state.chat_sessions = {}
        st.session_state.active_chat = None
        st.session_state.messages = []
        st.session_state.keyword_found = True
        st.session_state.suggested_prompts = []
        st.session_state.pending_prompts = []
        st.session_state.active_stream = None
        st.rerun()
    st.markdown("---")
    
    # --- MODEL & UPLOAD SETTINGS ---
    st.markdown("""
    <div style="display: flex; align-items: center; margin-bottom: 20px;">
        <div style="
            display: flex; 
            align-items: center; 
            justify-content: center; 
            height: 50px; 
            width: 50px; 
            border-radius: 8px; 
            background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
            color: white; 
            font-size: 24px; 
            margin-right: 12px;
            border: 2px solid #d4af37;
            box-shadow: 0 4px 8px rgba(0,0,0,0.2);
        ">⚖️</div>
        <div>
            <h3 style="margin: 0; color: #1e3c72; font-weight: bold;">Advocate AI Pro</h3>
            <p style="margin: 0; font-size: 12px; color: #666;">Legal Research Assistant</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.session_state.selected_model = st.selectbox("Select Model", [
        "openai/gpt-5", "anthropic/claude-sonnet-4", "google/gemini-2.5-pro",
        "x-ai/grok-code-fast-1", "google/gemini-2.5-flash", 
        "deepseek/deepseek-r1-distill-llama-70b:free", "openai/gpt-5-mini", "tngtech/deepseek-r1t2-chimera:free"
    ], index=0)
    
    st.success(f"Knowledge Base loaded with {folder_index.ntotal} chunks", icon="✅")
    
    if PYPDF2_AVAILABLE:
        uploaded_pdfs = st.file_uploader("Upload PDFs to use as context", type=["pdf"], accept_multiple_files=True)
        if uploaded_pdfs:
            st.session_state.uploaded_pdfs_data = uploaded_pdfs
            st.info(f"{len(uploaded_pdfs)} PDF(s) loaded. They will be used as context.")
        elif st.session_state.uploaded_pdfs_data:
            st.session_state.uploaded_pdfs_data = None
            
    st.markdown("---")

# ==========================
# MAIN CHAT INTERFACE
# ==========================
st.markdown(f"""
<div style="display: flex; align-items: center; margin-bottom: 20px; padding: 20px; background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%); border-radius: 10px; border-left: 5px solid #d4af37;">
    <div style="
        display: flex; 
        align-items: center; 
        justify-content: center; 
        height: 70px; 
        width: 70px; 
        border-radius: 10px; 
        background: white; 
        color: #1e3c72; 
        font-size: 32px; 
        margin-right: 20px;
        border: 3px solid #d4af37;
        box-shadow: 0 6px 12px rgba(0,0,0,0.15);
    ">⚖️</div>
    <div>
        <h1 style="margin: 0; color: white; font-weight: 700; font-size: 2.5rem;">Legal Research Assistant</h1>
        <p style="margin: 5px 0 0 0; color: #f0f2f6; font-size: 1.1rem; font-weight: 400;">
            Professional Legal Analysis & Case Research
        </p>
    </div>
</div>
""", unsafe_allow_html=True)

if st.session_state.active_chat and st.session_state.active_chat in st.session_state.chat_sessions:
    current_chat_name = st.session_state.chat_sessions[st.session_state.active_chat]["name"]
    st.subheader(f"💬 {current_chat_name}")

keyword_terms = st.session_state.get("keyword_highlight_input", "")

chat_container = st.container()
with chat_container:
    # Display chat history
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            content = message["content"]
            
            if message["role"] == "assistant":
                js_safe_content = content.replace("'", "\\'").replace('\n', ' ')
                st.markdown(
                    f'<button onclick="speakText(\'{js_safe_content}\')" class="stButton secondary">🔊 Read Aloud</button>',
                    unsafe_allow_html=True
                )

            # Render highlights
            if keyword_terms and keyword_terms.strip():
                highlighted_content, _ = highlight_with_style(content, keyword_terms)
                st.markdown(highlighted_content, unsafe_allow_html=True)
            else:
                st.markdown(content)
            
            if message.get("sources"):
                with st.expander("Show Sources (from Database)"):
                    for source in message["sources"]:
                        llm_summary = summarize_source_with_llm(source["text"], model=st.session_state.selected_model) 
                        st.markdown(f"**Source:** `{source['source']}`")
                        st.markdown(f"**Snippet (3 Lines):** {llm_summary}")
    
    # Handle active stream
    if st.session_state.active_stream:
        with st.chat_message("assistant"):
            # Stream the response live into the chat bubble
            full_response = st.write_stream(st.session_state.active_stream)
        
        # Stream is finished, now clean up
        
        # 1. Add the completed message to session state
        assistant_message = {
            "role": "assistant",
            "content": full_response,
            "sources": st.session_state.stream_sources
        }
        st.session_state.messages.append(assistant_message)
        
        # 2. Clear the stream states
        st.session_state.active_stream = None
        st.session_state.stream_sources = []
        
        # 3. Generate follow-up suggestions
        try:
            suggested_prompts = generate_suggested_prompts(st.session_state.last_user_prompt, full_response)
            st.session_state.suggested_prompts = suggested_prompts
        except Exception as e:
            st.session_state.suggested_prompts = []

        # 4. Auto-name the chat (if needed)
        if st.session_state.active_chat:
            current_chat = st.session_state.chat_sessions[st.session_state.active_chat]
            if current_chat["name"] == "New Chat":
                first_user_msg = st.session_state.last_user_prompt
                if first_user_msg:
                    new_chat_name = first_user_msg[:40].strip()
                    if len(first_user_msg) > 40:
                        new_chat_name += "..."
                    st.session_state.chat_sessions[st.session_state.active_chat]["name"] = new_chat_name
            
            # 5. Save messages to chat session
            st.session_state.chat_sessions[st.session_state.active_chat]["messages"] = st.session_state.messages

        # 6. Rerun to display the new suggestions and clear the stream
        st.rerun()

# ==================================
# OPTIMIZED PROMPT WORKFLOW
# ==================================

def handle_selected_prompt(selected_prompt):
    """Optimized version: sets up the stream generator in session state."""
    st.session_state.pending_prompts = []
    st.session_state.processing = True # Show processing indicator while RAG runs

    # RAG from Database
    all_sources = []
    try:
        chunks, sources = retrieve_and_rerank(selected_prompt, folder_index, folder_metadata, embed_model, reranker_model, top_k=3)
        all_sources.extend(sources)
    except Exception as e:
        pass  # Silent error handling

    unique_sources = {}
    for source in all_sources:
        key = source['source'] + source['text'][:100]
        if key not in unique_sources:
            unique_sources[key] = source
    final_rag_context = "\n\n".join([truncate_text(s['text'], max_words=200) for s in unique_sources.values()])
    
    # RAG from Uploaded PDFs
    uploaded_context_summary = ""
    if st.session_state.uploaded_pdfs_data and PYPDF2_AVAILABLE:
        full_pdf_text = extract_text_from_pdfs(st.session_state.uploaded_pdfs_data)
        if full_pdf_text:
            pdf_chunks = full_pdf_text.split('\n\n')
            pdf_metadata = [{'text': chunk, 'source': 'Uploaded Document'} for chunk in pdf_chunks if chunk.strip()]
            if pdf_metadata:
                pdf_search_results = bm25_search(selected_prompt, pdf_metadata, top_k=5)
                uploaded_context_summary = "\n\n".join([res['text'] for res in pdf_search_results])

    # Construct Final Prompt
    final_prompt = construct_final_prompt(selected_prompt, final_rag_context, uploaded_context_summary)
    
    # Create the generator and store it
    st.session_state.active_stream = call_openrouter(
        st.session_state.selected_model, 
        final_prompt, 
        list(unique_sources.values())
    )
    # Store the sources to be added to the message *after* streaming
    st.session_state.stream_sources = list(unique_sources.values())
    st.session_state.last_user_prompt = selected_prompt
    
    # We are no longer "processing" (waiting for RAG), we are about to stream.
    st.session_state.processing = False


# ==================================
# IMPROVED PROMPT SELECTION UI
# ==================================

if st.session_state.processing:
    st.markdown("""
    <div class="processing-indicator">
        ⚖️ Analyzing Your Legal Research Query
        <div class="processing-subtext">Reviewing case law, statutes, and relevant precedents...</div>
    </div>
    """, unsafe_allow_html=True)

elif st.session_state.pending_prompts:
    st.markdown('<div class="section-header">📋 Select Your Research Query Format</div>', unsafe_allow_html=True)
    st.markdown("Choose the most appropriate framing for your legal research question:")
    
    st.markdown('<div class="prompt-selection-container">', unsafe_allow_html=True)
    
    labels = [
        ("🎯", "Original Question", "Your initial query as stated"),
        ("📚", "Detailed Analysis", "Comprehensive legal framework exploration"),
        ("⚖️", "Case Law & Precedents", "Judicial interpretation and statutory references")
    ]
    
    col1, col2, col3 = st.columns(3)
    columns = [col1, col2, col3]
    
    for i, (prompt_text, (icon, label, desc), col) in enumerate(zip(st.session_state.pending_prompts, labels, columns)):
        with col:
            card_class = "original" if i == 0 else "refined"
            st.markdown(f"""
            <div class="prompt-card {card_class}">
                <div class="prompt-label">
                    <span class="icon">{icon}</span>
                    <span>{label}</span>
                </div>
                <div class="prompt-text">{prompt_text}</div>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button(f"Select Query {i+1}", key=f"pending_{i}", use_container_width=True, type="primary" if i == 0 else "secondary"):
                handle_selected_prompt(prompt_text)
                st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

elif st.session_state.suggested_prompts:
    st.markdown('<div class="section-header">💡 Suggested Follow-up Questions</div>', unsafe_allow_html=True)
    st.markdown("Click on any question below to continue your research:")
    
    for i, prompt_text in enumerate(st.session_state.suggested_prompts):
        suggestion_container = st.container()
        with suggestion_container:
            st.markdown(f'<div class="suggestion-container">', unsafe_allow_html=True)
            
            button_label = f"🔍 {prompt_text}"
            if st.button(
                button_label, 
                key=f"suggestion_{i}", 
                use_container_width=True,
                help="Click to explore this question"
            ):
                # Clear suggestions and generate new variations
                st.session_state.suggested_prompts = []
                
                if not st.session_state.active_chat:
                    chat_id = f"chat_{len(st.session_state.chat_sessions) + 1}_{random.randint(1000,9999)}"
                    new_chat_name = prompt_text[:30].strip() + "..." if len(prompt_text) > 30 else prompt_text.strip()
                    st.session_state.chat_sessions[chat_id] = {"name": new_chat_name, "messages": [], "created_at": datetime.now()}
                    st.session_state.active_chat = chat_id
                    st.session_state.messages = st.session_state.chat_sessions[chat_id]["messages"]
                
                st.session_state.messages.append({"role": "user", "content": prompt_text})
                variations = dynamic_query_generator(prompt_text)
                st.session_state.pending_prompts = variations
                
                if st.session_state.active_chat:
                    st.session_state.chat_sessions[st.session_state.active_chat]["messages"] = st.session_state.messages
                
                st.rerun()
            
            st.markdown('</div>', unsafe_allow_html=True)

# --- Streamlit Chat Input ---
chat_input_disabled = bool(st.session_state.pending_prompts) or st.session_state.processing or bool(st.session_state.active_stream)

if prompt := st.chat_input("Ask a legal question (e.g., 'What is the doctrine of Res Gestae?')", disabled=chat_input_disabled):
    st.session_state.suggested_prompts = []
    
    if not st.session_state.active_chat:
        chat_id = f"chat_{len(st.session_state.chat_sessions) + 1}_{random.randint(1000,9999)}"
        new_chat_name = prompt[:30].strip() + "..." if len(prompt) > 30 else prompt.strip()
        st.session_state.chat_sessions[chat_id] = {"name": new_chat_name, "messages": [], "created_at": datetime.now()}
        st.session_state.active_chat = chat_id
        st.session_state.messages = st.session_state.chat_sessions[chat_id]["messages"]
    
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    variations = dynamic_query_generator(prompt)
    st.session_state.pending_prompts = variations
    
    if st.session_state.active_chat:
        st.session_state.chat_sessions[st.session_state.active_chat]["messages"] = st.session_state.messages
    
    st.rerun()
